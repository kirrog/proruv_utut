import pandas as pd
from docx import Document
from docx.shared import Inches, Cm, RGBColor, Pt
from datetime import datetime

# Функция для добавления изображения в документ
def add_image_to_doc(doc, image_path, width=Inches(2.0)):
    if image_path:
        doc.add_picture(image_path, width=width)
        doc.paragraphs[-1].alignment = 1  # Центрирование изображения

with open('report_data.csv', 'r', encoding='windows-1251') as file:
    lines = file.readlines()

doc = Document()

style = doc.styles.add_style('CustomStyle', 1)
style.font.name = 'Times New Roman'
style.font.size = Pt(14)

table = doc.add_table(rows=1, cols=1)
table.autofit = False
table.columns[0].width = Cm(14.6)  # Ширина страницы A4
table.rows[0].height = Cm(29.7)  # Высота страницы A4
cell = table.cell(0, 0)
cell.vertical_alignment = 1  # Вертикальное центрирование

p = cell.add_paragraph()
p.alignment = 1  # Горизонтальное центрирование
p.style = 'CustomStyle'
p.add_run('Отчет по дефектам/поломкам на ноутбуках').bold = True
p.runs[0].font.color.rgb = RGBColor(0, 0, 0)

p = cell.add_paragraph()
p.alignment = 1  # Горизонтальное центрирование
p.style = 'CustomStyle'
p.add_run('Автор: ФИО').bold = True
p.runs[0].font.color.rgb = RGBColor(0, 0, 0)

p = cell.add_paragraph()
p.alignment = 1  # Горизонтальное центрирование
p.style = 'CustomStyle'
p.add_run(f'Дата: {datetime.now().strftime("%d.%m.%Y")}').bold = True
p.runs[0].font.color.rgb = RGBColor(0, 0, 0)

flag = 0
count = 1

for line in lines:
    if ':' in line:
        key, value = line.strip().split(':', 1)
        if key == "Заказчик":
            p = doc.add_heading('Информация о ноутбуке:', level=2)
            p.style = 'CustomStyle'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            p.runs[0].bold = True

            p = doc.add_paragraph(f'Заказчик: {value}')
            p.style = 'CustomStyle'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        if key == 'Модель ноутбука':
            p = doc.add_paragraph(f'Модель ноутбука: {value}')
            p.style = 'CustomStyle'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        elif key == 'Серийный номер':
            p = doc.add_paragraph(f'Серийный номер: {value}')
            p.style = 'CustomStyle'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        elif key == 'Изображение ноутбука':
            add_image_to_doc(doc, value)
            p = doc.add_paragraph()
        elif key == 'Дефекты/поломки':
            if flag == 0:
                p = doc.add_heading('Дефекты/поломки:', level=2)
                p.style = 'CustomStyle'
                p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                p.runs[0].bold = True
                flag = 1

            p = doc.add_paragraph(f'Описание дефекта: {value}')
            p.style = 'CustomStyle'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        elif key == 'Количество дефектов':
            p = doc.add_paragraph(f'Количество дефектов: {value}')
            p.style = 'CustomStyle'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            p = doc.add_paragraph()
        elif key == 'Рекомендации':
            p = doc.add_heading('Рекомендации:', level=2)
            p.style = 'CustomStyle'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            p.runs[0].bold = True

            p = doc.add_paragraph(f'Рекомендуемые действия: {value}')
            p.style = 'CustomStyle'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        elif key == 'Срочность ремонта':
            p = doc.add_paragraph(f'Срочность ремонта: {value}')
            p.style = 'CustomStyle'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        elif key == 'Стоимость ремонта':
            p = doc.add_paragraph(f'Стоимость ремонта: {value}')
            p.style = 'CustomStyle'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        elif key == 'Комментарии':
            p = doc.add_paragraph()
            p = doc.add_paragraph(f'Комментарии:')
            p.style = 'CustomStyle'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            p.runs[0].bold = True

            p = doc.add_paragraph(value)
            p.style = 'CustomStyle'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)

doc.save('report.docx')
